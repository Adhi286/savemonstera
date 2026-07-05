from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
import pandas as pd
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from database import supabase
from routers.auth import get_current_user

router = APIRouter()

@router.get("/{format}")
def export_data(format: str, user_id: str = Depends(get_current_user)):
    try:
        response = supabase.table("expenses").select("*").eq("user_id", user_id).order("date", desc=True).execute()
        data = response.data if response.data else []
        
        # Calculate summary
        total_spent = sum([ex["amount"] for ex in data if ex.get("type") == "debit"])
        highest_spend = max([ex["amount"] for ex in data if ex.get("type") == "debit"], default=0)
        lowest_spend = min([ex["amount"] for ex in data if ex.get("type") == "debit"], default=0)
        
        cat_totals = {}
        for ex in data:
            if ex.get("type") == "debit":
                cat = ex.get("category", "Other")
                cat_totals[cat] = cat_totals.get(cat, 0) + ex["amount"]
        sorted_cats = sorted(cat_totals.items(), key=lambda x: x[1], reverse=True)
        
        if format == "xlsx":
            return export_excel(data, total_spent, highest_spend, lowest_spend, sorted_cats)
        elif format == "pdf":
            return export_pdf(data, total_spent, highest_spend, lowest_spend, sorted_cats)
        elif format == "docx":
            return export_docx(data, total_spent, highest_spend, lowest_spend, sorted_cats)
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def export_excel(data, total, high, low, cats):
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='openpyxl')
    
    # Summary sheet
    summary_data = {
        "Metric": ["Total Spent (₹)", "Highest Spend (₹)", "Lowest Spend (₹)"] + [f"Category: {c[0]}" for c in cats],
        "Value": [total, high, low] + [c[1] for c in cats]
    }
    pd.DataFrame(summary_data).to_excel(writer, sheet_name="Summary", index=False)
    
    # Data sheet
    df = pd.DataFrame(data)
    if not df.empty:
        # keep specific columns
        df = df[["date", "description", "category", "amount", "mode", "status"]]
        df["amount"] = df["amount"].apply(lambda x: f"₹{x}")
    df.to_excel(writer, sheet_name="Transactions", index=False)
    
    writer.close()
    output.seek(0)
    
    from fastapi.responses import Response
    return Response(
        content=output.getvalue(), 
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=expenses.xlsx"}
    )

def export_pdf(data, total, high, low, cats):
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    c.drawString(100, 750, "SaveMonstera Expense Report")
    c.drawString(100, 730, f"Total Spent: ₹{total}")
    c.drawString(100, 715, f"Highest Spend: ₹{high}")
    c.drawString(100, 700, f"Lowest Spend: ₹{low}")
    
    y = 670
    c.drawString(100, y, "Category Rankings:")
    for cat in cats:
        y -= 15
        c.drawString(120, y, f"{cat[0]}: ₹{cat[1]}")
    
    y -= 30
    c.drawString(100, y, "Recent Transactions:")
    for ex in data[:20]: # show first 20 for pdf simplicity
        y -= 15
        if y < 50:
            c.showPage()
            y = 750
        c.drawString(100, y, f"{ex['date']} - {ex.get('description', '')[:20]} - ₹{ex['amount']}")
        
    c.save()
    output.seek(0)
    from fastapi.responses import Response
    return Response(
        content=output.getvalue(), 
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=expenses.pdf"}
    )

def export_docx(data, total, high, low, cats):
    doc = Document()
    doc.add_heading('SaveMonstera Expense Report', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f'Total Spent: ₹{total}')
    doc.add_paragraph(f'Highest Spend: ₹{high}')
    doc.add_paragraph(f'Lowest Spend: ₹{low}')
    
    doc.add_heading('Category Rankings', level=2)
    for cat in cats:
        doc.add_paragraph(f"{cat[0]}: ₹{cat[1]}", style='List Bullet')
        
    doc.add_heading('Transactions', level=1)
    if data:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Category'
        hdr_cells[3].text = 'Amount'
        
        for ex in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(ex['date'])
            row_cells[1].text = str(ex.get('description', ''))
            row_cells[2].text = str(ex.get('category', ''))
            row_cells[3].text = f"₹{ex['amount']}"
            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    from fastapi.responses import Response
    return Response(
        content=output.getvalue(), 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=expenses.docx"}
    )
