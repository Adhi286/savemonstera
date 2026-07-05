from docx import Document
import os

doc = Document()
doc.add_heading("SaveMonstera - Project Description", 0)

doc.add_heading("💡 Inspiration", level=1)
doc.add_paragraph("University life is stressful enough without constantly worrying about where your money went. We noticed that traditional budgeting apps are often too rigid, filled with financial jargon, and require tedious manual entry. Students need something fast, friendly, and smart. That’s why we built SaveMonstera—a financial companion that uses AI to make tracking expenses as easy as sending a text message.")

doc.add_heading("⚙️ What it does", level=1)
doc.add_paragraph("SaveMonstera is a web-based budget tracker designed specifically for the student lifestyle in India.")
doc.add_paragraph("• AI Expense Logging: Instead of filling out complex forms, students can just type naturally (e.g., \"Spent ₹450 on pizza for the study group\"). Our Gemini AI agent automatically parses the text, extracts the exact rupee amount, categorizes the expense as \"Food\", and logs the date.")
doc.add_paragraph("• Smart Budget Thresholds: The AI has autonomous logic built-in. If an expense keeps you under your weekly budget, it's auto-approved. If you're spending too fast, the agent flags it for review to help you stay accountable.")
doc.add_paragraph("• Beautiful Dashboard: A clean, modern, and mobile-friendly UI featuring a white and yellow aesthetic. It provides a quick glance at weekly and monthly remaining budgets in INR (₹), recent activity, and visual category breakdowns.")

doc.add_heading("🛠️ How we built it", level=1)
doc.add_paragraph("• Frontend: Built with HTML, CSS, and Vanilla JavaScript, styled with a modern, glass-like UI to ensure it feels premium without the heavy overhead of large frameworks.")
doc.add_paragraph("• Backend: A fast, asynchronous REST API powered by FastAPI (Python).")
doc.add_paragraph("• Database & Auth: We integrated Supabase for secure user authentication (JWT) and persistent PostgreSQL database storage for expenses and budgets.")
doc.add_paragraph("• AI Agent Engine: We utilized the Google Gemini AI SDK (gemini-2.5-flash) to create the natural language processing nodes that parse inputs and run our autonomous budget-checking logic.")
doc.add_paragraph("• Deployment: Containerized via Docker and deployed seamlessly using Google Cloud Agent Runtime / Cloud Run.")

doc.add_heading("🚧 Challenges we ran into", level=1)
doc.add_paragraph("Integrating the AI SDK into our strict backend environment threw us a few curveballs! We had to carefully manage dependency versions between legacy packages and the newest SDK releases to ensure the Gemini model could be invoked without crashing. We also had to implement a clever token-management system on the frontend to prevent infinite redirect loops during authentication.")

doc.add_heading("🏆 Accomplishments that we're proud of", level=1)
doc.add_paragraph("We successfully bridged the gap between a robust relational database (Supabase) and a generative AI model (Gemini). Watching the AI successfully parse a messy, colloquial sentence into a perfectly structured JSON database row for the very first time was an incredible feeling!")

doc.add_heading("📖 What we learned", level=1)
doc.add_paragraph("We learned a massive amount about integrating LLMs into structured software pipelines. We learned that prompt engineering is just as much about formatting output (forcing strictly structured JSON) as it is about guiding the AI's logic. We also deepened our knowledge of FastAPI middleware, Supabase JWT auth flows, and Docker containerization.")

doc.add_heading("🚀 What's next for SaveMonstera", level=1)
doc.add_paragraph("• Receipt Scanning: Adding a feature to upload photos of receipts and bills, using Gemini's multimodal vision capabilities to extract line items automatically.")
doc.add_paragraph("• Predictive Budgeting: Using historical spending data to warn students when they are statistically likely to run out of money before the end of the semester.")
doc.add_paragraph("• Shared Expenses: Allowing roommates or hostel-mates to split utility and grocery bills seamlessly within the app.")

output_path = "/home/adhibalan/.gemini/antigravity/brain/1c36d75a-6f13-426e-a66c-5a4ee96f3350/SaveMonstera_Project_Description.docx"
doc.save(output_path)
print(f"DOCX_CREATED_AT_{output_path}")
